from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Inches

from quizmark.models import MediaReference, QuizData


def export_docx(quiz: QuizData) -> bytes:
    doc = Document()
    doc.add_heading(quiz.title, level=1)

    if quiz.media:
        _add_media_block(doc, quiz.media, prefix="Quiz media")

    for idx, question in enumerate(quiz.questions, start=1):
        header = f"{idx}. {question.text}"
        if question.points is not None:
            header += f" ({question.points} pts)"
        doc.add_heading(header, level=2)

        question_media = _combined_media(question.media, question.image)
        if question_media:
            _add_media_block(doc, question_media, prefix="Question media")

        for answer in question.answers:
            mark = " (correct)" if answer.correct else ""
            text = answer.text or ""
            doc.add_paragraph(f"{answer.label}. {text}{mark}")

            answer_media = _combined_media(answer.media, answer.image)
            if answer_media:
                _add_media_block(doc, answer_media, prefix="Answer media")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_media_block(doc: Document, media_list: list[MediaReference], prefix: str) -> None:
    doc.add_paragraph(prefix + ":")
    for media in media_list:
        if media.kind == "image":
            if _try_add_image(doc, media.value):
                continue
            doc.add_paragraph(f"- image (missing): {media.value}")
            continue
        if media.kind == "audio":
            doc.add_paragraph(f"- audio: {media.value}")
            continue
        if media.kind == "video":
            doc.add_paragraph(f"- video: {media.value}")
            continue
        if media.kind == "attachment":
            doc.add_paragraph(f"- attachment: {media.value}")
            continue
        if media.kind == "math":
            doc.add_paragraph(f"- math: {media.value}")
            continue
        doc.add_paragraph(f"- {media.kind}: {media.value}")


def _combined_media(media_list, image) -> list[MediaReference]:
    items = list(media_list)
    if image and not any(m.kind == "image" and m.value == image.path for m in items):
        items.append(MediaReference(kind="image", value=image.path, line=image.line))
    return items


def _try_add_image(doc: Document, value: str) -> bool:
    path = Path(value)
    if not path.exists():
        return False
    try:
        doc.add_picture(str(path), width=Inches(5.5))
    except Exception:
        return False
    return True
